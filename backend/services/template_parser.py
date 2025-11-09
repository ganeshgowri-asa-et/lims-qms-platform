"""
Template Parser Service
Parses Excel and Word templates to extract form field definitions
"""
import openpyxl
from docx import Document
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import json


class TemplateParser:
    """Parse Excel and Word templates to generate form field definitions"""

    def __init__(self):
        self.field_type_mapping = {
            'text': 'TEXT',
            'number': 'NUMBER',
            'date': 'DATE',
            'datetime': 'DATETIME',
            'dropdown': 'DROPDOWN',
            'multiselect': 'MULTISELECT',
            'checkbox': 'CHECKBOX',
            'radio': 'RADIO',
            'file': 'FILE',
            'signature': 'SIGNATURE',
            'table': 'TABLE',
            'section': 'SECTION',
            'calculated': 'CALCULATED'
        }

    def parse_excel_template(self, file_path: str) -> Dict[str, Any]:
        """
        Parse Excel template to extract form fields

        Expected format:
        - Sheet 1: Form metadata (name, description, category)
        - Sheet 2: Field definitions (columns: field_name, field_label, field_type, is_required, options, validation_rules)
        """
        wb = openpyxl.load_workbook(file_path, data_only=True)

        template_data = {
            'name': '',
            'code': '',
            'description': '',
            'category': '',
            'fields': [],
            'validation_rules': [],
            'layout_config': {}
        }

        # Parse metadata sheet
        if 'Metadata' in wb.sheetnames or 'metadata' in wb.sheetnames:
            meta_sheet = wb['Metadata'] if 'Metadata' in wb.sheetnames else wb['metadata']
            template_data.update(self._parse_metadata_sheet(meta_sheet))

        # Parse fields sheet
        if 'Fields' in wb.sheetnames or 'fields' in wb.sheetnames:
            fields_sheet = wb['Fields'] if 'Fields' in wb.sheetnames else wb['fields']
            template_data['fields'] = self._parse_fields_sheet(fields_sheet)

        # Auto-detect fields from main sheet if no Fields sheet
        if not template_data['fields'] and wb.active:
            template_data['fields'] = self._auto_detect_fields_from_excel(wb.active)

        wb.close()
        return template_data

    def _parse_metadata_sheet(self, sheet) -> Dict[str, str]:
        """Parse metadata sheet with key-value pairs"""
        metadata = {}
        for row in sheet.iter_rows(min_row=1, values_only=True):
            if row[0] and row[1]:
                key = str(row[0]).lower().strip()
                value = str(row[1]).strip()
                if key in ['name', 'code', 'description', 'category', 'version']:
                    metadata[key] = value
        return metadata

    def _parse_fields_sheet(self, sheet) -> List[Dict[str, Any]]:
        """Parse fields definition sheet"""
        fields = []
        headers = []

        for idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if idx == 0:
                # First row is header
                headers = [str(cell).lower().strip() if cell else '' for cell in row]
                continue

            if not row[0]:  # Skip empty rows
                continue

            field_data = {}
            for i, cell in enumerate(row):
                if i < len(headers) and headers[i]:
                    field_data[headers[i]] = cell

            # Process field
            field = self._process_field_data(field_data, idx)
            if field:
                fields.append(field)

        return fields

    def _auto_detect_fields_from_excel(self, sheet) -> List[Dict[str, Any]]:
        """Auto-detect form fields from Excel template structure"""
        fields = []
        order = 0

        for row_idx, row in enumerate(sheet.iter_rows(min_row=1), start=1):
            for col_idx, cell in enumerate(row, start=1):
                if cell.value and isinstance(cell.value, str):
                    # Look for field labels (cells ending with ":", "*", etc.)
                    label = str(cell.value).strip()

                    if self._is_field_label(label):
                        field_name = self._generate_field_name(label)
                        field_type = self._detect_field_type(label, cell)

                        field = {
                            'field_name': field_name,
                            'field_label': label.rstrip(':*'),
                            'field_type': field_type,
                            'is_required': '*' in label or 'required' in label.lower(),
                            'order': order,
                            'metadata': {
                                'row': row_idx,
                                'column': col_idx,
                                'cell_address': cell.coordinate
                            }
                        }

                        fields.append(field)
                        order += 1

        return fields

    def _is_field_label(self, text: str) -> bool:
        """Check if text looks like a field label"""
        if len(text) < 2 or len(text) > 200:
            return False

        # Common patterns for field labels
        patterns = [
            r'^[\w\s]+:$',  # Ends with colon
            r'^\*[\w\s]+',  # Starts with asterisk
            r'^[\w\s]+\*$',  # Ends with asterisk
        ]

        return any(re.match(pattern, text.strip()) for pattern in patterns)

    def _generate_field_name(self, label: str) -> str:
        """Generate field name from label"""
        # Remove special characters and convert to snake_case
        name = re.sub(r'[^\w\s]', '', label)
        name = re.sub(r'\s+', '_', name.strip())
        return name.lower()

    def _detect_field_type(self, label: str, cell) -> str:
        """Detect field type from label and cell properties"""
        label_lower = label.lower()

        # Keyword-based detection
        if any(kw in label_lower for kw in ['signature', 'sign']):
            return 'SIGNATURE'
        elif any(kw in label_lower for kw in ['date', 'time']):
            if 'time' in label_lower:
                return 'DATETIME'
            return 'DATE'
        elif any(kw in label_lower for kw in ['number', 'quantity', 'amount', 'count']):
            return 'NUMBER'
        elif any(kw in label_lower for kw in ['file', 'attachment', 'upload']):
            return 'FILE'
        elif any(kw in label_lower for kw in ['dropdown', 'select', 'choose']):
            return 'DROPDOWN'
        elif any(kw in label_lower for kw in ['checkbox', 'check']):
            return 'CHECKBOX'

        # Default to text
        return 'TEXT'

    def _process_field_data(self, field_data: Dict[str, Any], order: int) -> Optional[Dict[str, Any]]:
        """Process field data from row"""
        if not field_data.get('field_name'):
            return None

        field = {
            'field_name': str(field_data.get('field_name', '')).strip(),
            'field_label': str(field_data.get('field_label', field_data.get('field_name', ''))).strip(),
            'field_type': self._normalize_field_type(field_data.get('field_type', 'TEXT')),
            'is_required': self._parse_boolean(field_data.get('is_required', False)),
            'is_readonly': self._parse_boolean(field_data.get('is_readonly', False)),
            'order': order,
            'section': field_data.get('section'),
            'help_text': field_data.get('help_text'),
            'placeholder': field_data.get('placeholder'),
            'default_value': field_data.get('default_value'),
        }

        # Parse options for dropdown/multiselect
        if field_data.get('options'):
            options_str = str(field_data['options'])
            field['options'] = [opt.strip() for opt in options_str.split(',')]

        # Parse validation rules
        if field_data.get('validation_rules'):
            field['validation_rules'] = self._parse_validation_rules(field_data['validation_rules'])

        return field

    def _normalize_field_type(self, field_type: str) -> str:
        """Normalize field type to standard values"""
        field_type_lower = str(field_type).lower().strip()
        return self.field_type_mapping.get(field_type_lower, 'TEXT')

    def _parse_boolean(self, value: Any) -> bool:
        """Parse boolean value from various formats"""
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.lower() in ['true', 'yes', '1', 'y']
        if isinstance(value, (int, float)):
            return bool(value)
        return False

    def _parse_validation_rules(self, rules_str: str) -> Dict[str, Any]:
        """Parse validation rules from string"""
        try:
            # Try JSON format
            return json.loads(rules_str)
        except:
            # Parse simple format: "min:0,max:100,pattern:^[0-9]+$"
            rules = {}
            for rule in str(rules_str).split(','):
                if ':' in rule:
                    key, value = rule.split(':', 1)
                    key = key.strip()
                    value = value.strip()

                    # Try to convert to appropriate type
                    try:
                        value = int(value)
                    except:
                        try:
                            value = float(value)
                        except:
                            pass

                    rules[key] = value
            return rules

    def parse_word_template(self, file_path: str) -> Dict[str, Any]:
        """
        Parse Word template to extract form fields

        Looks for:
        - Form fields (text boxes, dropdowns)
        - Tables with labeled rows
        - Text patterns indicating fields
        """
        doc = Document(file_path)

        template_data = {
            'name': Path(file_path).stem,
            'code': '',
            'description': '',
            'category': '',
            'fields': [],
            'layout_config': {}
        }

        # Extract metadata from document properties
        template_data['name'] = doc.core_properties.title or Path(file_path).stem
        template_data['description'] = doc.core_properties.subject or ''

        # Parse paragraphs for field patterns
        order = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if self._is_field_label(text):
                field = {
                    'field_name': self._generate_field_name(text),
                    'field_label': text.rstrip(':*'),
                    'field_type': self._detect_field_type(text, None),
                    'is_required': '*' in text or 'required' in text.lower(),
                    'order': order
                }
                template_data['fields'].append(field)
                order += 1

        # Parse tables
        for table in doc.tables:
            table_fields = self._parse_word_table(table, order)
            template_data['fields'].extend(table_fields)
            order += len(table_fields)

        return template_data

    def _parse_word_table(self, table, start_order: int) -> List[Dict[str, Any]]:
        """Parse table in Word document"""
        fields = []
        order = start_order

        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                label = cells[0].text.strip()
                if self._is_field_label(label):
                    field = {
                        'field_name': self._generate_field_name(label),
                        'field_label': label.rstrip(':*'),
                        'field_type': self._detect_field_type(label, None),
                        'is_required': '*' in label,
                        'order': order
                    }
                    fields.append(field)
                    order += 1

        return fields

    def parse_pdf_template(self, file_path: str) -> Dict[str, Any]:
        """
        Parse PDF template (basic extraction)
        Note: PDF parsing is limited - recommend using Excel/Word templates
        """
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        template_data = {
            'name': Path(file_path).stem,
            'code': '',
            'description': '',
            'category': '',
            'fields': [],
            'layout_config': {}
        }

        # Extract text from all pages
        order = 0
        for page in reader.pages:
            text = page.extract_text()
            lines = text.split('\n')

            for line in lines:
                line = line.strip()
                if self._is_field_label(line):
                    field = {
                        'field_name': self._generate_field_name(line),
                        'field_label': line.rstrip(':*'),
                        'field_type': self._detect_field_type(line, None),
                        'is_required': '*' in line,
                        'order': order
                    }
                    template_data['fields'].append(field)
                    order += 1

        return template_data

    def generate_template_code(self, name: str, category: str = None) -> str:
        """Generate unique template code"""
        # Extract initials from name
        words = name.split()
        initials = ''.join([w[0].upper() for w in words if w])[:4]

        # Add category prefix if available
        prefix = ''
        if category:
            cat_words = category.split()
            prefix = ''.join([w[0].upper() for w in cat_words if w])[:2] + '-'

        # Generate code
        code = f"{prefix}{initials}"
        return code
